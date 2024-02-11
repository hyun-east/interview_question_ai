import pandas as pd
import openai
import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
import time

window = tk.Tk()
window.title("AI 면접 질문 추출기")
window.geometry("640x360")
window.resizable(False, False)


n = 1
global tx
tx = "디렉토리 이름"

def keymake():
     key = entry1.get()
     f = open(file="API Key.txt", mode="w", encoding="utf-8")
     f.write(key)
     f.close()
     tk.messagebox.showinfo("안내", "API 키가 저장되었습니다.")

def keyinfo():
    tk.messagebox.showinfo("안내", "API KEY 발급 방법:\n사이트 접속 (https://platform.openai.com) 후에 회원가입을 해준다\n2. API 발급 메뉴로 들어 간다\n오른쪽 상단 메뉴의 View API Keys를 눌러준다\n3. 왼쪽 메뉴에서 Billing을 눌러준다\n결제 카드 등록\n4. API Keys를 눌러준다\nCreat new secret key")
def namesave():
    global name
    name = entry.get()
    label2 = tk.Label(window, text=name+"...Name Saved")
    label2.place(x=120, y=100, width=400, height=30)
def press():
    window.file = filedialog.askopenfile(
        title='파일 선택창',
        filetypes=(('xlsx files', '*.xlsx'), ('all files', '*.*')))
    print(window.file.name)
    tx = window.file.name
    label = tk.Label(window, text=tx)
    label.place(x=0, y=190, width=640, height=30)


def start_progress():
    dir_path = filedialog.askdirectory(parent=window, initialdir="/", title='저장 위치 선택')
    print(dir_path)
    time.sleep(0.05)
    tk.messagebox.showinfo("알림", "저장 위치가 선택되었습니다. 본 알림을 끄면 작업이 시작됩니다. ")
    f = open("API Key.txt", "r", encoding="utf-8")
    openai.api_key = f.read()
    df = pd.read_excel(window.file.name)
    print(df)
    title = df['title'].values.tolist()
    contents = df['contents'].values.tolist()
    global n
    n = len(title)
    results = []
    for j in range(n):
        print('=========================' + title[j] + ' 질문 (%d/%d)=========================' % (j + 1, n))
        completion = openai.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "user",
                 "content": '이 글을 기반으로 면접 질문을 추출해줘. 동기, 과정, 방법 등에 대해서 구체적이고 자세하게 물어보고, 기술적인 면, 수과학적 개념, 사회적, 인문학적인 부분 등에 대해서 깊이 있게 물어봐줘. 표면적으로 있는 내용을 물어보지 말고, 꼭 심층적인 내용으로 깊이있게 물어봐줘. 구체적일수록 좋아. 글에 나와있는 내용의 정의나 기본 개념도 물어봐줘. 면접 질문은 15개를 생성해 주면 돼. 대답은 다른 말 하지 말고 질문만 알려줘.존댓말을 사용하지 말고, 반말로 물어봐줘. 말투는 -요 형식으로 하지 말고 인가? 하라. 설명하시오, 하시오, 하는가? 했는가?  형식으로 작성해줘. 면접자를 지칭할 때는 면접자, 지원자 등의 어휘를 사용해줘.' +
                            contents[j]}
            ]
        )
        prog_bar['value']=j+1
        time.sleep(0.05)
        prog_bar.update()

        results.append(completion.choices[0].message.content)
        print(completion.choices[0].message.content)
        print('\n')


    con = ''
    for i in range(n):
        con += '\n\n'
        con += title[i]
        con += '\n\n'
        con += results[i]

    from docx import Document


    doc = Document()

    doc.save('%s/%s.docx' %(dir_path, name))

    doc = Document('%s/%s.docx'%(dir_path, name))

    doc.add_heading('%s님의 생활기록부 기반 면접 질문 Report' % name, level=0)

    doc.add_paragraph(
        '*인공지능에 의해 자동으로 생성되어서 어색한 표현이 있을 수 있습니다. \n*이상하거나 중복되는 질문을 고려해서 15개를 만들었으므로 모든 질문이 중요한 것은 아닐 수 있습니다.\n*모든 질문에 대해서 답변하려 하기보다 어느 정도 중요해 보이는 문제에 대해서 답변할 것을 추천합니다.')

    doc.add_paragraph(con)

    doc.save('%s/%s.docx' %(dir_path, name))

    prog_bar['value'] = 0
    tk.messagebox.showinfo("알림", "질문 추출이 완료되었습니다.")
    prog_bar.update()



entry1 = tk.Entry(window)
entry1.place(x=20, y=11, width=500, height=23)
entry = tk.Entry(window)
entry.place(x=120, y=70, width=200, height=30)
ttk.Button(window, text="Select Name", command=namesave).place(x=320, y=70, width=200, height=30)

ttk.Button(window, text="Select File", command=press).place(x=120, y=130, width=400, height=60)
ttk.Button(window, text="Start Process", command=start_progress).place(x=120, y=220, width=400, height=60)
ttk.Button(window, text="API Key 저장", command=keymake).place(x=530, y=10, width=100, height=25)
ttk.Button(window, text="key 발급 방법", command=keyinfo).place(x=530, y=300, width=100, height=25)




prog_bar = ttk.Progressbar(window, length=400, maximum=n+1)
prog_bar.place(x=120, y=300)

label4 = tk.Label(window, text='github.io/hyun-east  khd306090@gmail.com')
label4.place(x=0, y=330, width=640, height=30)

label4 = tk.Label(window, text='처음 사용하거나 API 키가 변경된 경우 API 키를 입력하고 저장해야 합니다. \' \" \' 등의 문자 없이 키만 입력하십시오.')
label4.place(x=0, y=40, width=640, height=25)

window.mainloop()
